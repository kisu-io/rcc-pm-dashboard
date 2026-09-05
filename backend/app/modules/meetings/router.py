# DDC-CWICR-OE: DataDrivenConstruction · OpenConstructionERP
# Copyright (c) 2026 Artem Boiko / DataDrivenConstruction
"""Meetings API routes.

Endpoints:
    GET    /                              - List meetings for a project
    POST   /                              - Create meeting (auto-generates meeting_number)
    POST   /import-summary                - Import meeting from transcript file (AI-powered)
    GET    /{meeting_id}                  - Get single meeting
    PATCH  /{meeting_id}                  - Update meeting
    DELETE /{meeting_id}                  - Delete meeting
    POST   /{meeting_id}/complete         - Mark meeting as completed
    GET    /{meeting_id}/calendar.ics     - Export meeting (or series) as an iCalendar feed
    GET    /{meeting_id}/export/pdf       - Export meeting minutes as PDF
"""

import io
import logging
import re
import uuid
from datetime import UTC, datetime

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from fastapi.responses import Response, StreamingResponse

from app.core.content_disposition import attachment_disposition
from app.dependencies import CurrentUserId, RequirePermission, SessionDep, verify_project_access
from app.modules.meetings import ics as ics_builder
from app.modules.meetings.schemas import (
    MEETING_TYPES,
    ActionItemEntry,
    ActionRegisterItemCreate,
    ActionRegisterItemResponse,
    ActionRegisterItemUpdate,
    AgendaItemEntry,
    AttendanceRow,
    AttendeeEntry,
    CheckInRequest,
    ExternalAttendeeRequest,
    ImportPreviewResponse,
    MaterializeRequest,
    MeetingActionsResponse,
    MeetingCreate,
    MeetingResponse,
    MeetingSeriesCreate,
    MeetingSeriesResponse,
    MeetingStatsResponse,
    MeetingUpdate,
    MinutesDistributeResponse,
    MinutesGenerateRequest,
    MinutesResponse,
    MinutesUpdate,
    OpenActionItemResponse,
    SeriesActionRegisterResponse,
)
from app.modules.meetings.service import MeetingService

router = APIRouter(tags=["meetings"])
logger = logging.getLogger(__name__)


def _get_service(session: SessionDep) -> MeetingService:
    return MeetingService(session)


def _meeting_to_response(meeting: object) -> MeetingResponse:
    """Build a MeetingResponse from a Meeting ORM object."""
    return MeetingResponse(
        id=meeting.id,  # type: ignore[attr-defined]
        project_id=meeting.project_id,  # type: ignore[attr-defined]
        meeting_number=meeting.meeting_number,  # type: ignore[attr-defined]
        meeting_type=meeting.meeting_type,  # type: ignore[attr-defined]
        title=meeting.title,  # type: ignore[attr-defined]
        meeting_date=meeting.meeting_date,  # type: ignore[attr-defined]
        location=meeting.location,  # type: ignore[attr-defined]
        chairperson_id=(
            str(meeting.chairperson_id) if meeting.chairperson_id else None  # type: ignore[attr-defined]
        ),
        attendees=meeting.attendees or [],  # type: ignore[attr-defined]
        agenda_items=meeting.agenda_items or [],  # type: ignore[attr-defined]
        action_items=meeting.action_items or [],  # type: ignore[attr-defined]
        minutes=meeting.minutes,  # type: ignore[attr-defined]
        status=meeting.status,  # type: ignore[attr-defined]
        document_ids=[str(d) for d in (getattr(meeting, "document_ids", None) or [])],
        series_id=getattr(meeting, "series_id", None),
        is_series_master=bool(getattr(meeting, "is_series_master", False)),
        created_by=meeting.created_by,  # type: ignore[attr-defined]
        metadata=getattr(meeting, "metadata_", {}),
        created_at=meeting.created_at,  # type: ignore[attr-defined]
        updated_at=meeting.updated_at,  # type: ignore[attr-defined]
    )


# ── List ──────────────────────────────────────────────────────────────────────


@router.get(
    "/",
    response_model=list[MeetingResponse],
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def list_meetings(
    session: SessionDep,
    project_id: uuid.UUID = Query(...),
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    offset: int = Query(default=0, ge=0),
    limit: int = Query(default=50, ge=1, le=100),
    type_filter: str | None = Query(default=None, alias="type"),
    status_filter: str | None = Query(default=None, alias="status"),
    search: str | None = Query(
        default=None,
        max_length=200,
        description="Free-text search across title, agenda, minutes, and meeting number.",
    ),
    sort_by: str | None = Query(default=None, description="Sort field: date, meeting_date, status, created_at"),
    sort_order: str = Query(default="desc", pattern="^(asc|desc)$"),
    service: MeetingService = Depends(_get_service),
) -> list[MeetingResponse]:
    """List meetings for a project with optional filters and search."""
    await verify_project_access(project_id, user_id, session)
    # Map friendly alias "date" to the actual model column
    resolved_sort = "meeting_date" if sort_by == "date" else sort_by
    meetings, _ = await service.list_meetings(
        project_id,
        offset=offset,
        limit=limit,
        meeting_type=type_filter,
        status_filter=status_filter,
        search=search,
        sort_by=resolved_sort,
        sort_order=sort_order,
    )
    return [_meeting_to_response(m) for m in meetings]


# ── Stats ────────────────────────────────────────────────────────────────────


@router.get(
    "/stats/",
    response_model=MeetingStatsResponse,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def meeting_stats(
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    project_id: uuid.UUID = Query(...),
    service: MeetingService = Depends(_get_service),
) -> MeetingStatsResponse:
    """Aggregate meeting statistics for a project.

    Returns total, breakdown by status and type, count of open action items
    across all meetings, and the next upcoming meeting date.
    """
    await verify_project_access(project_id, user_id, session)
    return await service.get_stats(project_id)


# ── Open Action Items ────────────────────────────────────────────────────────


@router.get(
    "/open-actions/",
    response_model=list[OpenActionItemResponse],
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def open_action_items(
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    project_id: uuid.UUID = Query(...),
    service: MeetingService = Depends(_get_service),
) -> list[OpenActionItemResponse]:
    """All open action items across all meetings in a project.

    Returns each action item with its parent meeting context (number, title, date).
    """
    await verify_project_access(project_id, user_id, session)
    return await service.get_open_actions(project_id)


# ── Create ────────────────────────────────────────────────────────────────────


@router.post("/", response_model=MeetingResponse, status_code=201)
async def create_meeting(
    data: MeetingCreate,
    user_id: CurrentUserId,
    session: SessionDep,
    _perm: None = Depends(RequirePermission("meetings.create")),
    service: MeetingService = Depends(_get_service),
) -> MeetingResponse:
    """Create a new meeting with auto-generated meeting number."""
    await verify_project_access(data.project_id, user_id, session)
    meeting = await service.create_meeting(data, user_id=user_id)
    return _meeting_to_response(meeting)


# ── Import Summary (AI-powered transcript parsing) ──────────────────────────


def _parse_vtt_transcript(content: str) -> list[dict[str, str]]:
    """Parse WebVTT (.vtt) transcript into structured segments.

    VTT format:
        WEBVTT

        00:00:00.000 --> 00:00:05.000
        Speaker Name: Hello everyone...

    Returns:
        List of dicts with 'speaker', 'text', and 'timestamp' keys.
    """
    segments: list[dict[str, str]] = []
    lines = content.strip().splitlines()
    current_speaker = ""
    current_text = ""
    current_ts = ""

    ts_pattern = re.compile(r"(\d{2}:\d{2}:\d{2}[\.,]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[\.,]\d{3})")

    for line in lines:
        line = line.strip()
        if not line or line.upper() == "WEBVTT" or line.startswith("NOTE"):
            continue

        ts_match = ts_pattern.match(line)
        if ts_match:
            # Save previous segment
            if current_text:
                segments.append({"speaker": current_speaker, "text": current_text.strip(), "timestamp": current_ts})
                current_text = ""
            current_ts = ts_match.group(1)
            continue

        # Skip numeric cue identifiers
        if line.isdigit():
            continue

        # Check for speaker tag: "Speaker Name: text"
        speaker_match = re.match(r"^<v\s+([^>]+)>(.*)$", line)
        if speaker_match:
            current_speaker = speaker_match.group(1).strip()
            current_text += " " + speaker_match.group(2).strip()
        elif ": " in line and len(line.split(": ", 1)[0]) < 50:
            parts = line.split(": ", 1)
            current_speaker = parts[0].strip()
            current_text += " " + parts[1].strip()
        else:
            current_text += " " + line

    # Final segment
    if current_text:
        segments.append({"speaker": current_speaker, "text": current_text.strip(), "timestamp": current_ts})

    return segments


def _parse_srt_transcript(content: str) -> list[dict[str, str]]:
    """Parse SRT subtitle format into structured segments.

    SRT format:
        1
        00:00:00,000 --> 00:00:05,000
        Speaker Name: Hello everyone...

    Returns:
        List of dicts with 'speaker', 'text', and 'timestamp' keys.
    """
    segments: list[dict[str, str]] = []
    blocks = re.split(r"\n\s*\n", content.strip())

    for block in blocks:
        lines = block.strip().splitlines()
        if len(lines) < 2:
            continue

        # Skip numeric index and timestamp line
        text_lines = []
        timestamp = ""
        for line in lines:
            if line.strip().isdigit():
                continue
            ts_match = re.match(
                r"(\d{2}:\d{2}:\d{2}[,\.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,\.]\d{3})",
                line.strip(),
            )
            if ts_match:
                timestamp = ts_match.group(1)
                continue
            text_lines.append(line.strip())

        text = " ".join(text_lines).strip()
        if not text:
            continue

        speaker = ""
        if ": " in text and len(text.split(": ", 1)[0]) < 50:
            parts = text.split(": ", 1)
            speaker = parts[0].strip()
            text = parts[1].strip()

        segments.append({"speaker": speaker, "text": text, "timestamp": timestamp})

    return segments


def _parse_plain_text(content: str) -> list[dict[str, str]]:
    """Parse plain text transcript (line by line).

    Attempts to detect speaker patterns like 'Name: text' or '[Name] text'.

    Returns:
        List of dicts with 'speaker', 'text', and 'timestamp' keys.
    """
    segments: list[dict[str, str]] = []
    for line in content.strip().splitlines():
        line = line.strip()
        if not line:
            continue

        speaker = ""
        text = line

        # Pattern: [Speaker Name] text
        bracket_match = re.match(r"^\[([^\]]+)\]\s*(.*)$", line)
        if bracket_match:
            speaker = bracket_match.group(1).strip()
            text = bracket_match.group(2).strip()
        # Pattern: Speaker Name: text
        elif ": " in line and len(line.split(": ", 1)[0]) < 50:
            parts = line.split(": ", 1)
            # Avoid splitting on time-like patterns (e.g., "10:30")
            if not re.match(r"^\d{1,2}$", parts[0]):
                speaker = parts[0].strip()
                text = parts[1].strip()

        if text:
            segments.append({"speaker": speaker, "text": text, "timestamp": ""})

    return segments


def _detect_source(filename: str, content: str) -> str:
    """Detect the meeting platform source from filename and content.

    Checks for known platform identifiers in the filename and in the first
    500 characters of the transcript content.

    Returns:
        Platform identifier: 'teams', 'google_meet', 'zoom', 'webex', or 'other'.
    """
    fn = filename.lower()
    ct = content[:500].lower()

    if "teams" in fn or "teams" in ct or "microsoft teams" in ct:
        return "teams"
    if "meet" in fn or "google meet" in ct or "meet.google" in ct:
        return "google_meet"
    if "zoom" in fn or "zoom" in ct or "zoom.us" in ct:
        return "zoom"
    if "webex" in fn or "cisco webex" in ct or "webex" in ct:
        return "webex"
    return "other"


def _infer_meeting_type(text: str) -> str:
    """Infer the meeting type from transcript content.

    Scans for domain-specific keywords to categorize the meeting.

    Every type in :data:`MEETING_TYPES` needs keywords here. A type the schema
    accepts but this function can never return is one the import path cannot
    reach at all, so the user has to correct the guess by hand on every
    transcript of that kind.

    Returns:
        One of :data:`MEETING_TYPES`, falling back to 'progress'.
    """
    lower = text[:5000].lower()
    safety_kw = [
        "safety",
        "incident",
        "hazard",
        "ppe",
        "osha",
        "near miss",
        "risk assessment",
        "toolbox talk",
        "jsa",
        "job safety",
    ]
    design_kw = [
        "design review",
        "architectural",
        "structural",
        "mep",
        "schematic",
        "drawing",
        "specification",
        "detail",
    ]
    subcontractor_kw = [
        "subcontractor",
        "sub-contractor",
        "trade",
        "bid",
        "quote",
        "scope of work",
        "sow",
    ]
    kickoff_kw = ["kickoff", "kick-off", "project start", "mobilization"]
    closeout_kw = ["closeout", "close-out", "handover", "deficiency", "punchlist", "punch list"]
    # Deliberately narrow, and deliberately tested last. A bare "cost" or
    # "budget" belongs to a progress meeting as often as to a commercial one,
    # so only terms that name the commercial process itself are listed here.
    commercial_kw = [
        "valuation",
        "payment application",
        "interim certificate",
        "final account",
        "cost report",
        "cost review",
        "commercial review",
        "quantity surveyor",
        "retention release",
    ]

    if any(kw in lower for kw in safety_kw):
        return "safety"
    if any(kw in lower for kw in design_kw):
        return "design"
    if any(kw in lower for kw in subcontractor_kw):
        return "subcontractor"
    if any(kw in lower for kw in kickoff_kw):
        return "kickoff"
    if any(kw in lower for kw in closeout_kw):
        return "closeout"
    if any(kw in lower for kw in commercial_kw):
        return "commercial"
    return "progress"


# Expanded keyword patterns for heuristic extraction

_ACTION_KEYWORDS = re.compile(
    r"\b("
    r"action\s*item|action\s*:?|todo\s*:?|to-do|task\s*:?|"
    r"will\s+do|need\s+to|needs\s+to|should|must|"
    r"deadline\s*:?|assigned\s+to|responsible\s*:?|"
    r"follow\s*-?\s*up|by\s+(monday|tuesday|wednesday|thursday|friday|"
    r"saturday|sunday|next\s+week|end\s+of\s+week|end\s+of\s+month|eow|eod|eom)|"
    r"please\s+\w+|make\s+sure|ensure\s+that|"
    r"i\s+will|we\s+will|let'?s\s+\w+|"
    r"can\s+you|could\s+you"
    r")\b",
    re.IGNORECASE,
)

_DECISION_KEYWORDS = re.compile(
    r"\b("
    r"decided|agreed|approved|confirmed|let'?s\s+go\s+with|"
    r"final\s+decision|conclusion\s*:?|resolution\s*:?|verdict\s*:?|"
    r"we\s+chose|the\s+plan\s+is|going\s+forward|"
    r"resolved|consensus|sign\s*-?\s*off|signed\s+off|"
    r"green\s*-?\s*light|go\s+ahead|proceed\s+with"
    r")\b",
    re.IGNORECASE,
)

_TOPIC_KEYWORDS = re.compile(
    r"\b(agenda|topic|item\s+\d|point\s+\d|discuss|discussion|update\s+on|report\s+on)\b",
    re.IGNORECASE,
)

_DUE_DATE_PATTERN = re.compile(
    r"\b(?:by|due|before|until)\s+"
    r"(\d{4}-\d{2}-\d{2}|"
    r"(?:monday|tuesday|wednesday|thursday|friday|saturday|sunday)|"
    r"(?:next\s+(?:week|month))|"
    r"(?:end\s+of\s+(?:week|month|day))|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{1,2})"
    r"\b",
    re.IGNORECASE,
)

# ``ActionItemEntry.due_date`` is locked to a strict ISO ^\d{4}-\d{2}-\d{2}$
# regex by pydantic.  The heuristic extractor however captures human
# phrases ("Friday", "next week", "end of month") and stuffs them into
# the same slot - which would 500 the endpoint on validation if we
# handed it through unfiltered.  This sentinel strips anything that
# isn't an ISO date; the original hint is parked on the action item's
# free-form description so reviewers can still see the speaker's
# intent.
_ISO_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")


def _coerce_iso_due_date(raw: str | None) -> str | None:
    """Return ``raw`` if it matches ISO YYYY-MM-DD, else ``None``.

    Used to scrub heuristic / LLM-extracted hints like ``"Friday"`` or
    ``"next week"`` that would otherwise trip ``ActionItemEntry``'s
    strict pattern validator and raise a 500 from the import endpoint.
    """
    if raw is None:
        return None
    return raw if _ISO_DATE_RE.match(str(raw).strip()) else None


def _extract_meeting_data_heuristic(
    segments: list[dict[str, str]],
    filename: str,
    raw_text: str = "",
) -> dict:
    """Extract meeting structure from transcript segments using heuristics.

    Identifies:
    - Attendees from speaker tags
    - Action items from keywords (expanded set including 'please', 'need to', 'let's')
    - Key decisions from keywords (expanded set including 'go ahead', 'sign off')
    - Meeting title from filename or first meaningful line
    - Meeting type from content keywords
    - Source platform from filename and content

    Args:
        segments: Parsed transcript segments with speaker, text, timestamp.
        filename: Original filename for title derivation and source detection.
        raw_text: Full raw text content for source detection.

    Returns:
        Dict with title, meeting_type, attendees, agenda_items, action_items,
        minutes, source, decisions, key_topics, and segments_count.
    """
    attendees: dict[str, str] = {}  # name -> role
    action_items: list[dict] = []
    decisions: list[dict] = []
    discussion_topics: list[str] = []
    all_text_parts: list[str] = []

    for seg in segments:
        speaker = seg.get("speaker", "").strip()
        text = seg.get("text", "").strip()

        if speaker and speaker not in attendees:
            attendees[speaker] = ""

        if text:
            all_text_parts.append(text)

        # Check for action items
        if _ACTION_KEYWORDS.search(text):
            owner = speaker or "TBD"
            # Try to extract a due date reference
            due_match = _DUE_DATE_PATTERN.search(text)
            due_hint = due_match.group(1) if due_match else None
            action_items.append(
                {
                    "description": text[:500],
                    "owner_id": None,
                    "owner_name": owner,
                    "due_date": due_hint,
                    "status": "open",
                }
            )

        # Check for decisions
        if _DECISION_KEYWORDS.search(text):
            decisions.append({"decision": text[:500], "made_by": speaker or ""})

        # Check for agenda/topic mentions
        if _TOPIC_KEYWORDS.search(text):
            discussion_topics.append(text[:300])

    # Derive title from filename
    title = "Imported Meeting"
    clean_name = re.sub(r"\.(txt|vtt|srt|docx|pdf)$", "", filename, flags=re.IGNORECASE)
    clean_name = clean_name.replace("_", " ").replace("-", " ").strip()
    if clean_name:
        title = clean_name

    # Build minutes from all text
    minutes_text = "\n".join(all_text_parts[:200])  # Cap at ~200 lines
    if decisions:
        minutes_text += "\n\n--- Key Decisions ---\n" + "\n".join(f"- {d['decision']}" for d in decisions[:20])

    # Build attendee list
    attendee_list = [{"name": name, "company": "", "role": "", "status": "present"} for name in attendees]

    # Build agenda items from discussion topics
    agenda_list = [{"topic": topic, "presenter": None, "notes": None} for topic in discussion_topics[:20]]

    # Detect source platform and meeting type
    full_text = raw_text or " ".join(all_text_parts[:100])
    source = _detect_source(filename, full_text)
    meeting_type = _infer_meeting_type(full_text)

    # Extract key topics (unique discussion topics, deduped)
    key_topics: list[str] = []
    seen_topics: set[str] = set()
    for topic in discussion_topics[:20]:
        short = topic[:100].strip()
        low = short.lower()
        if low not in seen_topics:
            seen_topics.add(low)
            key_topics.append(short)

    return {
        "title": title,
        "meeting_type": meeting_type,
        "attendees": attendee_list,
        "agenda_items": agenda_list,
        "action_items": action_items,
        "minutes": minutes_text[:10000],
        "source": source,
        "decisions": decisions[:20],
        "key_topics": key_topics[:15],
        "segments_count": len(segments),
    }


async def _extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text content from uploaded file based on extension.

    Supports: .txt, .vtt, .srt, .docx, .pdf
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("txt", "vtt", "srt"):
        # Try UTF-8 first, then latin-1 as fallback
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("latin-1", errors="replace")

    if ext == "docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed, falling back to raw text extraction")
            # Fallback: extract text from docx XML
            import zipfile

            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                if "word/document.xml" in zf.namelist():
                    xml_content = zf.read("word/document.xml").decode("utf-8", errors="replace")
                    # Strip XML tags to get plain text
                    return re.sub(r"<[^>]+>", " ", xml_content).strip()
            return ""

    if ext == "pdf":
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages[:50]:  # Cap at 50 pages
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("pdfplumber not installed, cannot extract PDF text")
            return ""
        except Exception as exc:
            logger.warning("Failed to extract text from PDF: %s", exc)
            return ""

    # Unknown format - try as plain text
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        return file_content.decode("latin-1", errors="replace")


_AI_MEETING_SYSTEM = (
    "You are a construction project meeting transcript analyzer. "
    "You extract structured meeting data from construction industry transcripts. "
    "Return valid JSON only. Be precise and extract only what is explicitly stated."
)

_AI_MEETING_PROMPT_TEMPLATE = """Analyze this meeting transcript and extract structured data.

TRANSCRIPT:
{transcript_text}

Extract the following in JSON format:
{{
  "title": "Meeting title (derive from main topic discussed)",
  "meeting_type": "__MEETING_TYPES__",
  "key_topics": ["topic 1", "topic 2"],
  "attendees": [
    {{"name": "Person Name", "company": "Company if mentioned", "role": "Role if mentioned"}}
  ],
  "action_items": [
    {{"description": "What needs to be done", "owner": "Person responsible", "due_date": "YYYY-MM-DD if mentioned, null otherwise"}}
  ],
  "decisions": [
    {{"decision": "What was decided", "made_by": "Who decided"}}
  ],
  "summary": "2-3 sentence summary of the meeting",
  "next_meeting": "Date/time if mentioned, null otherwise"
}}

Rules:
- Extract ONLY what is explicitly stated in the transcript
- For action items, identify keywords: "will do", "action", "deadline", "by Friday", "need to", "should"
- For decisions: "decided", "agreed", "approved", "confirmed", "let's go with"
- For attendees: look for speaker names before colons or in brackets
- meeting_type: infer from content (safety topics = safety, valuations/payment applications/cost reports = commercial, schedule and progress updates = progress, design/drawings = design, subcontractor/trade = subcontractor, etc.)
- For construction meetings: look for RFIs, submittals, change orders, schedule updates, safety incidents, trade coordination
"""  # noqa: E501

# The model can only answer with a type it was shown, so the alternation is
# built from the schema's vocabulary rather than typed out again. A word added
# to MEETING_TYPES and forgotten here would be a type the import path never
# proposes to anyone.
_AI_MEETING_PROMPT = _AI_MEETING_PROMPT_TEMPLATE.replace("__MEETING_TYPES__", "|".join(MEETING_TYPES))


async def _extract_with_ai(
    session: object,
    user_id: str,
    text_content: str,
    extracted: dict,
) -> bool:
    """Try to enhance extracted meeting data using AI.

    Calls the configured AI provider to analyze the transcript and merges results
    with the heuristic extraction (AI takes priority).

    Args:
        session: Database session for loading AI settings.
        user_id: Current user ID for looking up their AI API keys.
        text_content: Full transcript text (will be truncated for the AI prompt).
        extracted: Mutable dict of heuristic-extracted data to merge AI results into.

    Returns:
        True if AI was successfully used, False otherwise.
    """
    try:
        from sqlalchemy import select

        from app.modules.ai.ai_client import call_ai, extract_json, resolve_provider_key_model
        from app.modules.ai.models import AISettings

        result = await session.execute(  # type: ignore[union-attr]
            select(AISettings).where(AISettings.user_id == uuid.UUID(str(user_id)))
        )
        ai_settings = result.scalar_one_or_none()

        if not ai_settings:
            return False

        provider, api_key, model_override = resolve_provider_key_model(ai_settings)

        # Truncate transcript for AI prompt (leave room for the prompt template)
        transcript_preview = text_content[:8000]
        ai_prompt = _AI_MEETING_PROMPT.format(transcript_text=transcript_preview)

        raw_response, _tokens = await call_ai(
            provider=provider,
            api_key=api_key,
            system=_AI_MEETING_SYSTEM,
            prompt=ai_prompt,
            max_tokens=4096,
            model=model_override,
        )

        ai_data = extract_json(raw_response)
        if not isinstance(ai_data, dict):
            return False

        # Merge AI results with heuristic results (AI takes priority)
        if ai_data.get("title"):
            extracted["title"] = ai_data["title"]
        if ai_data.get("meeting_type") and ai_data["meeting_type"] in MEETING_TYPES:
            extracted["meeting_type"] = ai_data["meeting_type"]
        if ai_data.get("key_topics") and isinstance(ai_data["key_topics"], list):
            extracted["key_topics"] = [str(t)[:200] for t in ai_data["key_topics"][:15] if t]
        if ai_data.get("attendees") and isinstance(ai_data["attendees"], list):
            extracted["attendees"] = [
                {
                    "name": a.get("name", "Unknown"),
                    "company": a.get("company", ""),
                    "role": a.get("role", ""),
                    "status": "present",
                }
                for a in ai_data["attendees"]
                if isinstance(a, dict) and a.get("name")
            ]
        if ai_data.get("agenda_items") and isinstance(ai_data["agenda_items"], list):
            extracted["agenda_items"] = [
                {
                    "topic": item.get("topic", ""),
                    "presenter": item.get("presenter"),
                    "notes": item.get("notes"),
                }
                for item in ai_data["agenda_items"]
                if isinstance(item, dict) and item.get("topic")
            ]
        if ai_data.get("action_items") and isinstance(ai_data["action_items"], list):
            extracted["action_items"] = [
                {
                    "description": item.get("description", ""),
                    "owner_id": None,
                    "owner_name": item.get("owner", item.get("owner_name", "TBD")),
                    "due_date": item.get("due_date"),
                    "status": item.get("status", "open"),
                }
                for item in ai_data["action_items"]
                if isinstance(item, dict) and item.get("description")
            ]
        if ai_data.get("summary"):
            extracted["minutes"] = str(ai_data["summary"])[:10000]
        if ai_data.get("decisions") and isinstance(ai_data["decisions"], list):
            extracted["decisions"] = [
                (
                    {"decision": d.get("decision", str(d)), "made_by": d.get("made_by", "")}
                    if isinstance(d, dict)
                    else {"decision": str(d), "made_by": ""}
                )
                for d in ai_data["decisions"][:20]
            ]

        return True

    except Exception as exc:
        # AI is optional - log and continue with heuristic results
        logger.debug("AI-enhanced transcript parsing skipped: %s", exc)
        return False


def _build_preview_response(extracted: dict, ai_used: bool) -> ImportPreviewResponse:
    """Build an ImportPreviewResponse from the extracted data dict.

    Args:
        extracted: Dict of extracted meeting data (from heuristics and/or AI).
        ai_used: Whether AI enhancement was applied.

    Returns:
        ImportPreviewResponse with all extracted fields mapped.
    """
    from app.modules.meetings.schemas import (
        ImportPreviewActionItem,
        ImportPreviewAttendee,
        ImportPreviewDecision,
    )

    attendees = [
        ImportPreviewAttendee(
            name=a.get("name", "Unknown"),
            company=a.get("company", ""),
            role=a.get("role", ""),
        )
        for a in extracted.get("attendees", [])
    ]

    action_items = [
        ImportPreviewActionItem(
            description=a.get("description", ""),
            owner=a.get("owner_name", a.get("owner", "TBD")),
            due_date=a.get("due_date"),
        )
        for a in extracted.get("action_items", [])
        if a.get("description")
    ]

    raw_decisions = extracted.get("decisions", [])
    decisions = [
        ImportPreviewDecision(
            decision=d.get("decision", str(d)) if isinstance(d, dict) else str(d),
            made_by=d.get("made_by", "") if isinstance(d, dict) else "",
        )
        for d in raw_decisions
    ]

    return ImportPreviewResponse(
        title=extracted.get("title", "Imported Meeting"),
        meeting_type=extracted.get("meeting_type", "progress"),
        source=extracted.get("source", "other"),
        summary=extracted.get("minutes", "")[:2000],
        key_topics=extracted.get("key_topics", []),
        attendees=attendees,
        action_items=action_items,
        decisions=decisions,
        agenda_items=extracted.get("agenda_items", []),
        minutes=extracted.get("minutes", ""),
        ai_enhanced=ai_used,
        segments_parsed=extracted.get("segments_count", 0),
    )


@router.post("/import-summary/")
async def import_meeting_summary(
    session: SessionDep,
    project_id: uuid.UUID = Query(...),
    file: UploadFile = File(...),
    preview: bool = Query(default=False),
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.create")),
    service: MeetingService = Depends(_get_service),
) -> MeetingResponse | ImportPreviewResponse:
    """Import a meeting summary from a transcript file.

    Accepts: .txt, .vtt, .srt, .docx, .pdf files (transcripts/notes).

    Uses heuristic parsing to extract attendees, action items, decisions,
    and topics. When AI API keys are configured, the transcript is also
    sent to an LLM for higher-quality structured extraction.

    Query Parameters:
        preview: If true, returns extracted data without creating the meeting.
            The caller can then present the data for user review and call
            this endpoint again with preview=false to create the meeting.
    """
    # Project-ownership gate - pre-fix this endpoint accepted project_id
    # from the query string and went straight to file parsing / meeting
    # create with no cross-tenant verification, letting any user with the
    # meetings.create role inject meetings + documents into a foreign
    # project. ``verify_project_access`` returns 404 (not 403) so the
    # endpoint can't double as a UUID-existence oracle.
    await verify_project_access(project_id, str(user_id) if user_id else "", session)

    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    # Validate file extension
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    allowed_extensions = {"txt", "vtt", "srt", "docx", "pdf"}
    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format '.{ext}'. Accepted: {', '.join(sorted(allowed_extensions))}",
        )

    content = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail="File is empty")

    # Magic-byte sniff - filename is hostile-supplied. Reject any payload
    # whose first bytes don't match the declared format before we hand
    # the buffer to python-docx / PyMuPDF / VTT/SRT parsers. Plain-text
    # transcripts (txt/vtt/srt) have no canonical signature, so we use a
    # denylist of known-binary prefixes instead.
    head = content[:8]
    if ext == "pdf":
        # Allow a few bytes of leading whitespace/BOM (some scanners).
        stripped = head.lstrip(b"\x00\xef\xbb\xbf \t\r\n")
        if not stripped.startswith(b"%PDF-"):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File does not look like a valid PDF (missing %PDF- signature).",
            )
    elif ext == "docx":
        if not head.startswith(b"PK\x03\x04"):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File does not look like a valid .docx (missing ZIP signature).",
            )
    else:  # txt / vtt / srt - text-only formats
        for sig in (b"MZ", b"\x7fELF", b"\xca\xfe\xba\xbe", b"PK\x03\x04", b"\xd0\xcf\x11\xe0", b"%PDF-"):
            if head.startswith(sig):
                raise HTTPException(
                    status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                    detail="File does not look like a plain-text transcript (binary signature detected).",
                )

    # Extract text
    text_content = await _extract_text_from_file(content, file.filename)
    if not text_content.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract any text from the file. Please check the file format.",
        )

    # Parse into segments based on format
    if ext == "vtt":
        segments = _parse_vtt_transcript(text_content)
    elif ext == "srt":
        segments = _parse_srt_transcript(text_content)
    else:
        segments = _parse_plain_text(text_content)

    if not segments:
        raise HTTPException(
            status_code=400,
            detail="Could not parse any content from the transcript.",
        )

    # Extract meeting data using heuristics
    extracted = _extract_meeting_data_heuristic(segments, file.filename, text_content)

    # Try AI-enhanced parsing if available
    ai_used = await _extract_with_ai(service.session, str(user_id), text_content, extracted)

    # If preview mode, return extracted data without creating the meeting
    if preview:
        return _build_preview_response(extracted, ai_used)

    # Create meeting from extracted data
    today = datetime.now(UTC).strftime("%Y-%m-%d")

    attendees_data = [
        {
            "name": att.get("name", "Unknown"),
            "company": att.get("company", ""),
            "status": att.get("status", "present"),
        }
        for att in extracted.get("attendees", [])
    ]

    agenda_data = [
        {
            "number": str(idx + 1),
            "topic": item.get("topic", "Discussion item"),
            "presenter": item.get("presenter"),
            "notes": item.get("notes"),
        }
        for idx, item in enumerate(extracted.get("agenda_items", []))
    ]

    # Coerce free-form due_date hints ("Friday", "next week", ...) to
    # ``None`` so we don't trip ``ActionItemEntry``'s ISO-date regex
    # when materialising the action items below. The raw hint is folded
    # into the description so the reviewer can still see the intent.
    action_data: list[dict] = []
    for item in extracted.get("action_items", []):
        if not item.get("description"):
            continue
        raw_due = item.get("due_date")
        iso_due = _coerce_iso_due_date(raw_due)
        description = item.get("description", "")
        if raw_due and iso_due is None:
            # Surface the speaker's intent without breaking validation.
            description = f"{description} (due: {raw_due})"[:1000]
        action_data.append(
            {
                "description": description,
                "owner_id": item.get("owner_id"),
                "due_date": iso_due,
                "status": item.get("status", "open"),
            }
        )

    meeting_type = extracted.get("meeting_type", "progress")
    if meeting_type not in MEETING_TYPES:
        meeting_type = "progress"

    meeting_create = MeetingCreate(
        project_id=project_id,
        meeting_type=meeting_type,
        title=extracted.get("title", "Imported Meeting"),
        meeting_date=today,
        location=None,
        chairperson_id=None,
        attendees=[
            AttendeeEntry(
                name=att["name"],
                company=att.get("company"),
                status=att.get("status", "present"),
            )
            for att in attendees_data
        ],
        agenda_items=[
            AgendaItemEntry(
                number=item.get("number"),
                topic=item.get("topic", "Discussion item"),
                presenter=item.get("presenter"),
                notes=item.get("notes"),
            )
            for item in agenda_data
        ],
        action_items=[
            ActionItemEntry(
                description=item["description"],
                owner_id=item.get("owner_id"),
                due_date=item.get("due_date"),
                status=item.get("status", "open"),
            )
            for item in action_data
        ],
        minutes=extracted.get("minutes"),
        status="completed",
        metadata={
            "imported_from": file.filename,
            "import_source": extracted.get("source", "other"),
            "ai_enhanced": ai_used,
            "segments_parsed": extracted.get("segments_count", 0),
            "decisions": extracted.get("decisions", []),
            "key_topics": extracted.get("key_topics", []),
        },
    )

    meeting = await service.create_meeting(meeting_create, user_id=user_id)

    logger.info(
        "Meeting imported from transcript: file=%s, attendees=%d, actions=%d, ai=%s",
        file.filename,
        len(attendees_data),
        len(action_data),
        ai_used,
    )

    # Cross-link: save transcript and create Document record in
    # Documents hub.  Uses the ORM Document model directly so the row
    # picks up timestamps + defaults from the Base mixin and stays in
    # sync with any future schema migration.  Best-effort: a failure
    # here MUST NOT break the meeting create - the transcript file is
    # already on disk and the meeting itself is persisted.
    try:
        from pathlib import Path as _Path

        from app.modules.documents.models import Document

        upload_dir = _Path.home() / ".openestimator" / "uploads" / str(project_id)
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_uuid = uuid.uuid4().hex[:12]
        safe_name = re.sub(r"[^\w.\-]", "_", file.filename or "transcript")
        storage_name = f"{file_uuid}_{safe_name}"
        transcript_path = upload_dir / storage_name
        transcript_path.write_bytes(content)

        doc = Document(
            project_id=project_id,
            name=file.filename or "transcript",
            description=f"Meeting transcript: {extracted.get('title', '')}",
            category="correspondence",
            file_size=len(content),
            mime_type=file.content_type or "text/plain",
            file_path=str(transcript_path),
            version=1,
            uploaded_by=str(user_id) if user_id else "",
            tags=["meeting", "transcript"],
        )
        service.session.add(doc)
        await service.session.flush()
        logger.info("Cross-linked meeting transcript -> document %s", doc.id)
    except Exception:
        logger.exception("Failed to cross-link meeting transcript to Documents hub")

    return _meeting_to_response(meeting)


# ── Get ───────────────────────────────────────────────────────────────────────


@router.get(
    "/{meeting_id}",
    response_model=MeetingResponse,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def get_meeting(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> MeetingResponse:
    """Get a single meeting."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    return _meeting_to_response(meeting)


# ── Update ────────────────────────────────────────────────────────────────────


@router.patch("/{meeting_id}", response_model=MeetingResponse)
async def update_meeting(
    meeting_id: uuid.UUID,
    data: MeetingUpdate,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MeetingResponse:
    """Update a meeting."""
    existing = await service.get_meeting(meeting_id)
    await verify_project_access(existing.project_id, str(user_id), session)
    meeting = await service.update_meeting(meeting_id, data)
    return _meeting_to_response(meeting)


# ── Delete ────────────────────────────────────────────────────────────────────


@router.delete("/{meeting_id}", status_code=204)
async def delete_meeting(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.delete")),
    service: MeetingService = Depends(_get_service),
) -> None:
    """Delete a meeting."""
    existing = await service.get_meeting(meeting_id)
    await verify_project_access(existing.project_id, str(user_id), session)
    await service.delete_meeting(meeting_id)


# ── Complete ──────────────────────────────────────────────────────────────────


@router.post("/{meeting_id}/complete/", response_model=MeetingResponse)
async def complete_meeting(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MeetingResponse:
    """Mark a meeting as completed.

    Requires status to be 'scheduled' or 'in_progress'.
    Draft meetings must be scheduled first.
    Open action items are automatically converted to tasks.
    """
    existing = await service.get_meeting(meeting_id)
    await verify_project_access(existing.project_id, str(user_id), session)
    meeting = await service.complete_meeting(meeting_id, user_id=user_id)
    return _meeting_to_response(meeting)


# ── iCalendar (ICS) Export ────────────────────────────────────────────────────


def _meeting_ics_payload(meeting: object) -> dict:
    """Flatten an ORM Meeting into the plain mapping the pure ICS builder reads.

    Keeps ``app.modules.meetings.ics`` free of any ORM/DB knowledge - the pure
    builder only ever sees plain dicts, which is what its unit tests feed it too.
    """
    return {
        "id": meeting.id,  # type: ignore[attr-defined]
        "title": meeting.title,  # type: ignore[attr-defined]
        "meeting_number": meeting.meeting_number,  # type: ignore[attr-defined]
        "meeting_type": meeting.meeting_type,  # type: ignore[attr-defined]
        "meeting_date": meeting.meeting_date,  # type: ignore[attr-defined]
        "location": meeting.location,  # type: ignore[attr-defined]
        "chairperson_id": (
            str(meeting.chairperson_id) if meeting.chairperson_id else None  # type: ignore[attr-defined]
        ),
        "status": meeting.status,  # type: ignore[attr-defined]
        "attendees": meeting.attendees or [],  # type: ignore[attr-defined]
        "agenda_items": meeting.agenda_items or [],  # type: ignore[attr-defined]
        "minutes": meeting.minutes,  # type: ignore[attr-defined]
        "recurrence_rule": getattr(meeting, "recurrence_rule", None),
        "metadata": getattr(meeting, "metadata_", {}) or {},
    }


@router.get(
    "/{meeting_id}/calendar.ics",
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def export_meeting_ics(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> Response:
    """Export a meeting as an iCalendar (.ics) feed.

    Returns a single VEVENT, or - when the meeting is a recurring series master -
    the same VEVENT carrying an RRULE so a calendar client expands the whole
    series. The body is ``text/calendar`` and downloads as ``.ics`` for a
    one-click "add to calendar".
    """
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    ics_text = ics_builder.calendar_from_meeting(_meeting_ics_payload(meeting))
    filename = f"meeting_{meeting.meeting_number}.ics"
    return Response(
        content=ics_text,
        media_type=ics_builder.MEDIA_TYPE,
        headers={"Content-Disposition": attachment_disposition(filename)},
    )


# Export-convention mirror of the route above (matches ``/{meeting_id}/export/pdf/``).
@router.get(
    "/{meeting_id}/export/ics/",
    include_in_schema=False,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def export_meeting_ics_alias(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> Response:
    """Alias of ``/{meeting_id}/calendar.ics`` following the export-URL convention."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    ics_text = ics_builder.calendar_from_meeting(_meeting_ics_payload(meeting))
    filename = f"meeting_{meeting.meeting_number}.ics"
    return Response(
        content=ics_text,
        media_type=ics_builder.MEDIA_TYPE,
        headers={"Content-Disposition": attachment_disposition(filename)},
    )


# ── Recurring Series ─────────────────────────────────────────────────────────


def _attendance_to_row(att: object) -> AttendanceRow:
    """Build an AttendanceRow response from an ORM record."""
    return AttendanceRow(
        id=att.id,  # type: ignore[attr-defined]
        meeting_id=att.meeting_id,  # type: ignore[attr-defined]
        user_id=att.user_id,  # type: ignore[attr-defined]
        external_name=att.external_name,  # type: ignore[attr-defined]
        checked_in_at=att.checked_in_at,  # type: ignore[attr-defined]
        signature_image_path=att.signature_image_path,  # type: ignore[attr-defined]
        created_at=att.created_at,  # type: ignore[attr-defined]
        updated_at=att.updated_at,  # type: ignore[attr-defined]
    )


@router.post(
    "/series/",
    response_model=MeetingSeriesResponse,
    status_code=201,
)
async def create_meeting_series(
    data: MeetingSeriesCreate,
    user_id: CurrentUserId,
    session: SessionDep,
    _perm: None = Depends(RequirePermission("meetings.create")),
    service: MeetingService = Depends(_get_service),
) -> MeetingSeriesResponse:
    """Create a recurring meeting series (master + optional first occurrences).

    Body carries the master-meeting fields plus an RFC 5545 ``recurrence_rule``
    (``FREQ=WEEKLY;BYDAY=MO;COUNT=12``).  If ``materialize_until`` is set,
    occurrences are pre-created up to that date in the same call.
    """
    await verify_project_access(data.project_id, user_id, session)
    master, occurrences = await service.create_series(data, user_id=user_id)
    return MeetingSeriesResponse(
        series_id=master.id,
        master=_meeting_to_response(master),
        occurrences=[_meeting_to_response(o) for o in occurrences],
    )


@router.post(
    "/series/{master_id}/materialize/",
    response_model=MeetingSeriesResponse,
)
async def materialize_meeting_series(
    master_id: uuid.UUID,
    body: MaterializeRequest,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.create")),
    service: MeetingService = Depends(_get_service),
) -> MeetingSeriesResponse:
    """Materialise series occurrences up to the ``until`` date.

    Idempotent - already-existing occurrences are not duplicated.
    """
    master = await service.get_meeting(master_id)
    await verify_project_access(master.project_id, str(user_id), session)
    if not master.is_series_master:
        raise HTTPException(
            status_code=400,
            detail="Meeting is not a series master",
        )
    until_dt = datetime.strptime(body.until, "%Y-%m-%d").replace(tzinfo=UTC)
    new_occurrences = await service.generate_occurrences(
        str(master.id),
        until_dt,
        user_id=str(user_id) if user_id else None,
    )
    return MeetingSeriesResponse(
        series_id=master.id,
        master=_meeting_to_response(master),
        occurrences=[_meeting_to_response(o) for o in new_occurrences],
    )


# Underscore mirror - module convention is to expose hyphen + underscore.
@router.post("/series/{master_id}/materialize", include_in_schema=False)
async def materialize_meeting_series_no_slash(
    master_id: uuid.UUID,
    body: MaterializeRequest,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.create")),
    service: MeetingService = Depends(_get_service),
) -> MeetingSeriesResponse:
    return await materialize_meeting_series(  # type: ignore[return-value]
        master_id=master_id,
        body=body,
        session=session,
        user_id=user_id,
        _perm=None,
        service=service,
    )


# ── Attendance Check-in ──────────────────────────────────────────────────────


@router.post(
    "/{meeting_id}/check-in/",
    response_model=AttendanceRow,
)
async def check_in_attendee(
    meeting_id: uuid.UUID,
    body: CheckInRequest,
    session: SessionDep,
    user_id: CurrentUserId,
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> AttendanceRow:
    """Check the current user in for the given meeting.

    Idempotent: re-checking-in updates the existing row (single row per
    (meeting, user)).  Optional ``signature_image_data`` (base64 PNG or
    JPEG, or a ``data:image/...`` URL) is decoded and persisted on disk.
    """
    if not user_id:
        raise HTTPException(status_code=401, detail="Authentication required")
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.check_in(
        meeting_id,
        str(user_id),
        signature_image_data=body.signature_image_data,
    )
    return _attendance_to_row(row)


@router.post(
    "/{meeting_id}/external-attendee/",
    response_model=AttendanceRow,
    status_code=201,
)
async def record_external_attendee(
    meeting_id: uuid.UUID,
    body: ExternalAttendeeRequest,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> AttendanceRow:
    """Record a walk-in / non-system attendee by name only."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.record_external_attendee(
        meeting_id,
        body.name,
        signature_image_data=body.signature_image_data,
    )
    return _attendance_to_row(row)


@router.get(
    "/{meeting_id}/attendance/",
    response_model=list[AttendanceRow],
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def list_attendance(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> list[AttendanceRow]:
    """List all attendance rows for a meeting (system users + walk-ins)."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    rows = await service.list_attendance(meeting_id)
    return [_attendance_to_row(r) for r in rows]


# ── Action register (carry-over across a series) ─────────────────────────────


def _action_to_response(action: dict) -> ActionRegisterItemResponse:
    """Build an ActionRegisterItemResponse from a serialized action dict."""
    return ActionRegisterItemResponse.model_validate(action)


@router.get(
    "/{meeting_id}/actions/",
    response_model=MeetingActionsResponse,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def list_meeting_actions(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> MeetingActionsResponse:
    """List a meeting's own action items plus the ones brought forward into it.

    Brought-forward items are still-open actions raised in an earlier meeting of
    the same recurring series that have not yet been closed.
    """
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    own, brought = await service.list_meeting_actions(meeting)
    return MeetingActionsResponse(
        meeting_id=meeting_id,
        own=[_action_to_response(a) for a in own],
        brought_forward=[_action_to_response(a) for a in brought],
    )


@router.post(
    "/{meeting_id}/actions/",
    response_model=ActionRegisterItemResponse,
    status_code=201,
)
async def add_meeting_action(
    meeting_id: uuid.UUID,
    data: ActionRegisterItemCreate,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> ActionRegisterItemResponse:
    """Add a tracked action item (owner + due date required) to a meeting."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    action = await service.add_action(meeting, data, user_id=str(user_id) if user_id else None)
    return _action_to_response(action)


@router.patch("/actions/{action_id}", response_model=ActionRegisterItemResponse)
async def update_meeting_action(
    action_id: uuid.UUID,
    data: ActionRegisterItemUpdate,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> ActionRegisterItemResponse:
    """Update a tracked action item.

    Moving the status to ``done`` or ``cancelled`` closes the action for the
    whole series.
    """
    row = await service.get_action(action_id)
    await verify_project_access(row.project_id, str(user_id), session)
    action = await service.update_action(row, data)
    return _action_to_response(action)


@router.delete("/actions/{action_id}", status_code=204)
async def delete_meeting_action(
    action_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> None:
    """Delete a tracked action item."""
    row = await service.get_action(action_id)
    await verify_project_access(row.project_id, str(user_id), session)
    await service.delete_action(row)


@router.get(
    "/series/{master_id}/actions/",
    response_model=SeriesActionRegisterResponse,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def series_action_register(
    master_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> SeriesActionRegisterResponse:
    """The per-series action register - every action across the series.

    ``master_id`` is any meeting in the series (the series is keyed by its
    master id). Returns each action with its status plus a roll-up count.
    """
    master = await service.get_meeting(master_id)
    await verify_project_access(master.project_id, str(user_id), session)
    series_id, actions, summary = await service.series_action_register(master)
    return SeriesActionRegisterResponse(
        series_id=series_id,
        total=summary["total"],
        open=summary["open"],
        in_progress=summary["in_progress"],
        done=summary["done"],
        cancelled=summary["cancelled"],
        overdue=summary["overdue"],
        actions=[_action_to_response(a) for a in actions],
    )


# ── Auto-draft minutes ────────────────────────────────────────────────────────


@router.post(
    "/{meeting_id}/minutes/generate/",
    response_model=MinutesResponse,
)
async def generate_meeting_minutes(
    meeting_id: uuid.UUID,
    body: MinutesGenerateRequest,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MinutesResponse:
    """Generate (or refresh) the draft minutes for a meeting.

    Builds a structured draft: attendees present/absent, per-agenda discussion
    and decision, action items (brought-forward first) and the next meeting
    date. The draft is human-confirmed before it is issued.
    """
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.generate_minutes(meeting, body, user_id=str(user_id) if user_id else None)
    return MinutesResponse.model_validate(row)


@router.get(
    "/{meeting_id}/minutes/",
    response_model=MinutesResponse,
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def get_meeting_minutes(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> MinutesResponse:
    """Get the draft or issued minutes for a meeting."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.get_minutes_row(meeting_id)
    if row is None:
        raise HTTPException(status_code=404, detail="No minutes generated for this meeting yet")
    return MinutesResponse.model_validate(row)


@router.patch(
    "/{meeting_id}/minutes/",
    response_model=MinutesResponse,
)
async def update_meeting_minutes(
    meeting_id: uuid.UUID,
    data: MinutesUpdate,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MinutesResponse:
    """Apply human edits to a draft minutes document before it is issued."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.get_minutes_row(meeting_id)
    if row is None:
        raise HTTPException(status_code=404, detail="No minutes generated for this meeting yet")
    updated = await service.update_minutes(row, data)
    return MinutesResponse.model_validate(updated)


@router.post(
    "/{meeting_id}/minutes/issue/",
    response_model=MinutesResponse,
)
async def issue_meeting_minutes(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MinutesResponse:
    """Issue the minutes after the readiness validation passes."""
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.get_minutes_row(meeting_id)
    if row is None:
        raise HTTPException(status_code=404, detail="No minutes generated for this meeting yet")
    issued = await service.issue_minutes(row, user_id=str(user_id) if user_id else None)
    return MinutesResponse.model_validate(issued)


@router.post(
    "/{meeting_id}/minutes/distribute/",
    response_model=MinutesDistributeResponse,
)
async def distribute_meeting_minutes(
    meeting_id: uuid.UUID,
    session: SessionDep,
    user_id: CurrentUserId = None,  # type: ignore[assignment]
    _perm: None = Depends(RequirePermission("meetings.update")),
    service: MeetingService = Depends(_get_service),
) -> MinutesDistributeResponse:
    """Distribute the issued minutes to the meeting's attendees.

    Notifies every attendee linked to a real user via the platform notification
    service (in-app plus any configured connector).
    """
    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(user_id), session)
    row = await service.get_minutes_row(meeting_id)
    if row is None:
        raise HTTPException(status_code=404, detail="No minutes generated for this meeting yet")
    updated, notified = await service.distribute_minutes(meeting, row, user_id=str(user_id) if user_id else None)
    return MinutesDistributeResponse(
        minutes_id=updated.id,
        recipients=len(notified),
        notified_user_ids=notified,
    )


# ── PDF Export ───────────────────────────────────────────────────────────────


@router.get(
    "/{meeting_id}/export/pdf/",
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def export_meeting_pdf(
    meeting_id: uuid.UUID,
    session: SessionDep = None,  # type: ignore[assignment]
    _user: CurrentUserId = None,  # type: ignore[assignment]
) -> StreamingResponse:
    """Export meeting minutes as a PDF document."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from sqlalchemy import select

    from app.core.pdf_fonts import (
        BODY_FONT,
        BOLD_FONT,
        pdf_font_for_text,
        pdf_style_for_text,
        pdf_table_paragraph_rows,
        register_pdf_fonts,
    )
    from app.modules.meetings.models import Meeting
    from app.modules.projects.models import Project

    register_pdf_fonts()

    result = await session.execute(select(Meeting).where(Meeting.id == meeting_id))
    meeting = result.scalar_one_or_none()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    await verify_project_access(meeting.project_id, str(_user), session)

    # Fetch project name
    proj_result = await session.execute(select(Project.name).where(Project.id == meeting.project_id))
    project_name = proj_result.scalar_one_or_none() or "Unknown Project"

    # ── Build PDF ────────────────────────────────────────────────────────
    PAGE_WIDTH, PAGE_HEIGHT = A4
    MARGIN = 20 * mm
    USABLE_WIDTH = PAGE_WIDTH - 2 * MARGIN

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(
        "MeetingTitle",
        parent=styles["Normal"],
        fontName=BOLD_FONT,
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=4 * mm,
    )
    style_subtitle = ParagraphStyle(
        "MeetingSubtitle",
        parent=styles["Normal"],
        fontName=BODY_FONT,
        fontSize=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceAfter=6 * mm,
    )
    style_heading = ParagraphStyle(
        "SectionHeading",
        parent=styles["Normal"],
        fontName=BOLD_FONT,
        fontSize=12,
        spaceBefore=6 * mm,
        spaceAfter=3 * mm,
    )
    style_body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=BODY_FONT,
        fontSize=9,
        leading=12,
        alignment=TA_LEFT,
    )
    style_small = ParagraphStyle(
        "Small",
        parent=styles["Normal"],
        fontName=BODY_FONT,
        fontSize=8,
        textColor=colors.HexColor("#777777"),
    )
    style_cell = ParagraphStyle(
        "InfoCell",
        parent=styles["Normal"],
        fontName=BODY_FONT,
        fontSize=9,
        leading=11,
    )
    style_cell_label = ParagraphStyle("InfoLabel", parent=style_cell, fontName=BOLD_FONT)
    # The attendance and action tables named a face for their header row only,
    # so reportlab drew every body cell in its own default. Keeping Helvetica
    # here keeps those rows looking exactly as they did; pdf_style_for_text
    # still escalates a cell Helvetica cannot draw, which is what carried the
    # Chinese before.
    style_grid_cell = ParagraphStyle(
        "GridCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
    )
    style_grid_head = ParagraphStyle("GridHead", parent=style_grid_cell, fontName=BOLD_FONT)

    def _first_column_label(_row_index: int, col_index: int):
        """The label column is bold, every other column is a plain value."""
        return style_cell_label if col_index == 0 else None

    elements: list = []

    # Header
    elements.append(Paragraph("Meeting Minutes", style_title))
    elements.append(Paragraph(project_name, pdf_style_for_text(style_subtitle, project_name)))
    elements.append(Paragraph(meeting.title, pdf_style_for_text(style_heading, meeting.title)))

    # Meeting info table
    info_data = [
        ["Date:", meeting.meeting_date or "N/A"],
        ["Location:", meeting.location or "N/A"],
        ["Type:", (meeting.meeting_type or "").replace("_", " ").title()],
        ["Meeting #:", meeting.meeting_number],
        ["Status:", (meeting.status or "").replace("_", " ").title()],
    ]
    # Paragraph cells rather than bare strings. A bare cell is drawn through
    # canvas.drawString, which neither wraps nor shapes: a site address longer
    # than the column ran off the right hand side of the sheet, and a Thai or
    # Devanagari value was mis-arranged whatever face it was given. A Paragraph
    # does both and carries its own face and size, so the font commands and the
    # FONTNAME and FONTSIZE that used to sit here are gone rather than left to
    # describe a layout nothing obeys.
    info_data = pdf_table_paragraph_rows(info_data, style_cell, style_for=_first_column_label)
    info_table = Table(info_data, colWidths=[30 * mm, USABLE_WIDTH - 30 * mm])
    info_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    elements.append(info_table)
    elements.append(Spacer(1, 4 * mm))

    # Attendees
    attendees = meeting.attendees or []
    if attendees:
        elements.append(Paragraph("Attendees", style_heading))
        att_data = [["Name", "Company", "Status"]]
        for att in attendees:
            if isinstance(att, dict):
                att_data.append(
                    [
                        att.get("name", ""),
                        att.get("company", att.get("role", "")),
                        att.get("status", "").replace("_", " ").title(),
                    ]
                )
        # Paragraph cells for the reason given at the info table above, on the
        # same Helvetica body face reportlab was already drawing these rows in.
        att_data = pdf_table_paragraph_rows(att_data, style_grid_cell, header_style=style_grid_head, header_rows=1)
        att_table = Table(
            att_data,
            colWidths=[USABLE_WIDTH * 0.4, USABLE_WIDTH * 0.35, USABLE_WIDTH * 0.25],
        )
        att_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(att_table)

    # Agenda items
    agenda = meeting.agenda_items or []
    if agenda:
        elements.append(Paragraph("Agenda", style_heading))
        for idx, item in enumerate(agenda, 1):
            if isinstance(item, dict):
                topic = item.get("topic", item.get("title", ""))
                presenter = item.get("presenter", "")
                notes = item.get("notes", "")
                line = f"<b>{idx}.</b> {topic}"
                if presenter:
                    line += f"  <i>({presenter})</i>"
                elements.append(Paragraph(line, pdf_style_for_text(style_body, f"{topic}{presenter}")))
                if notes:
                    elements.append(
                        Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{notes}", pdf_style_for_text(style_small, notes))
                    )

    # Action items
    actions = meeting.action_items or []
    if actions:
        elements.append(Paragraph("Action Items", style_heading))
        act_data = [["#", "Description", "Owner", "Due Date", "Status"]]
        for idx, ai in enumerate(actions, 1):
            if isinstance(ai, dict):
                status_str = (
                    "Completed" if ai.get("completed") else (ai.get("status", "Open").replace("_", " ").title())
                )
                act_data.append(
                    [
                        str(idx),
                        ai.get("description", ""),
                        ai.get("owner", ai.get("owner_id", "")),
                        ai.get("due_date", ""),
                        status_str,
                    ]
                )
        # Paragraph cells for the reason given at the info table above. The
        # description is the column this matters most in: it is a sentence in
        # a 40 percent column and it was printed over the owner beside it.
        act_data = pdf_table_paragraph_rows(act_data, style_grid_cell, header_style=style_grid_head, header_rows=1)
        act_table = Table(
            act_data,
            colWidths=[
                USABLE_WIDTH * 0.06,
                USABLE_WIDTH * 0.40,
                USABLE_WIDTH * 0.20,
                USABLE_WIDTH * 0.17,
                USABLE_WIDTH * 0.17,
            ],
            # As on the minutes, and sooner here: this table is 9pt on 11pt
            # leading against the minutes' 8pt on 10pt and its description
            # column is narrower, so the same item ran out of frame at about
            # 2200 characters rather than 2900. Two geometries is why a length
            # cap would have been the wrong fix, since one number is wrong for
            # one of the two documents and the same meeting would render its
            # minutes and fail its export.
            splitInRow=1,
        )
        act_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(act_table)

    # Footer timestamp
    elements.append(Spacer(1, 10 * mm))
    generated_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    elements.append(Paragraph(f"Generated: {generated_at}", style_small))

    # Build document
    buf = io.BytesIO()

    def _header_footer(canvas_obj, doc):  # type: ignore[no-untyped-def]
        canvas_obj.saveState()
        canvas_obj.setFillColor(colors.HexColor("#999999"))
        running_head = f"{project_name} - {meeting.title}"
        canvas_obj.setFont(pdf_font_for_text(running_head), 7)
        canvas_obj.drawString(MARGIN, PAGE_HEIGHT - 12 * mm, running_head)
        canvas_obj.setFont(BODY_FONT, 7)
        canvas_obj.drawRightString(
            PAGE_WIDTH - MARGIN,
            10 * mm,
            f"Page {doc.page}",
        )
        canvas_obj.restoreState()

    frame = Frame(MARGIN, MARGIN, USABLE_WIDTH, PAGE_HEIGHT - 2 * MARGIN, id="main")
    doc = BaseDocTemplate(buf, pagesize=A4)
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=_header_footer)])
    doc.build(elements)

    buf.seek(0)
    safe_title = meeting.title.replace(" ", "_")[:50]
    filename = f"meeting_{meeting.meeting_number}_{safe_title}.pdf"

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={
            "Content-Disposition": attachment_disposition(filename),
            # The page furniture here is English literals and no locale reaches
            # this builder, so the route declares English rather than leaving
            # the Accept-Language middleware to answer a question about bytes it
            # never saw. The minutes themselves are whatever the attendees
            # typed, which is exactly why the labels around them should not
            # claim to be in the reader's language.
            "Content-Language": "en",
        },
    )


@router.get(
    "/{meeting_id}/minutes/export/pdf/",
    dependencies=[Depends(RequirePermission("meetings.read"))],
)
async def export_minutes_pdf(
    meeting_id: uuid.UUID,
    session: SessionDep = None,  # type: ignore[assignment]
    _user: CurrentUserId = None,  # type: ignore[assignment]
    service: MeetingService = Depends(_get_service),
) -> StreamingResponse:
    """Export the confirmed meeting minutes as a structured PDF.

    Renders the human-confirmed minutes document (present/absent, per-agenda
    discussion and decision, action items with brought-forward and overdue
    markers, next meeting date, summary) - not the raw meeting record.
    """
    from sqlalchemy import select

    from app.modules.meetings.pdf import build_minutes_pdf, minutes_pdf_filename
    from app.modules.projects.models import Project

    meeting = await service.get_meeting(meeting_id)
    await verify_project_access(meeting.project_id, str(_user), session)
    minutes = await service.get_minutes_row(meeting_id)
    if minutes is None:
        raise HTTPException(status_code=404, detail="No minutes generated for this meeting yet")

    proj_result = await session.execute(select(Project.name).where(Project.id == meeting.project_id))
    project_name = proj_result.scalar_one_or_none() or "Unknown Project"

    content: dict = minutes.content if isinstance(minutes.content, dict) else {}
    pdf_bytes = build_minutes_pdf(meeting, minutes, project_name)
    filename = minutes_pdf_filename(meeting, content)
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": attachment_disposition(filename),
            # English for the reason given on the meeting export above.
            "Content-Language": "en",
        },
    )
